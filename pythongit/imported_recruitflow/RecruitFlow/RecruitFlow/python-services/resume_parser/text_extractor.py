"""
Text Extraction Module
Handles extraction of text from various document formats (PDF, DOCX)
"""
import io
import logging
import os
from typing import Optional
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf_with_ocr(file_content: bytes) -> str:
    """
    Extract text from image-based PDF using OCR
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text as string
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        
        logger.info("Attempting OCR extraction for image-based PDF")
        
        # Convert PDF pages to images
        images = convert_from_bytes(file_content, dpi=300)
        
        text_content = []
        for i, image in enumerate(images, 1):
            logger.info(f"OCR processing page {i}/{len(images)}")
            # Extract text from image using Tesseract
            page_text = pytesseract.image_to_string(image)
            if page_text.strip():
                text_content.append(page_text)
        
        full_text = "\n".join(text_content)
        logger.info(f"OCR extracted {len(full_text)} chars from {len(images)} pages")
        return full_text.strip()
        
    except ImportError as e:
        logger.warning(f"OCR libraries not available: {e}")
        return ""
    except Exception as e:
        logger.error(f"OCR extraction failed: {str(e)}")
        return ""


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from a PDF file
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text as string
    """
    try:
        text_content = []
        
        # Create a BytesIO object from the file content
        pdf_file = io.BytesIO(file_content)
        
        # Use pdfplumber to extract text
        with pdfplumber.open(pdf_file) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        logger.info(f"Page {page_num}: extracted {len(page_text)} chars")
                    else:
                        logger.warning(f"Page {page_num}: no text extracted (might be image-based)")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
        
        # Join all pages with newlines
        full_text = "\n".join(text_content)
        logger.info(f"Total extracted text from PDF: {len(full_text)} chars")
        
        # If no text was extracted, try OCR
        if not full_text.strip():
            logger.info("No text extracted, attempting OCR...")
            return extract_text_from_pdf_with_ocr(file_content)
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF with pdfplumber: {str(e)}")
        # Try a more basic extraction approach if pdfplumber fails
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = []
            
            logger.info(f"Trying PyPDF2 extraction for {len(pdf_reader.pages)} pages")
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            full_text = "\n".join(text_content)
            logger.info(f"PyPDF2 extracted {len(full_text)} chars")
            return full_text.strip()
        except Exception as e2:
            logger.error(f"PyPDF2 also failed: {str(e2)}")
            raise Exception(f"Failed to extract text from PDF (tried pdfplumber and PyPDF2): {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from a DOCX file
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Extracted text as string
    """
    try:
        # Create a BytesIO object from the file content
        docx_file = io.BytesIO(file_content)
        
        # Use python-docx to extract text
        document = Document(docx_file)
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # Join all text with newlines
        full_text = "\n".join(text_content)
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text from a file based on its extension
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        if file_path.lower().endswith('.pdf'):
            return extract_text_from_pdf(file_content)
        elif file_path.lower().endswith(('.docx', '.doc')):
            return extract_text_from_docx(file_content)
        else:
            logger.error(f"Unsupported file format: {file_path}")
            return None
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return None